__author__ = "Tony Li"
__credits__ = ["Tony Li", "new guy"]
__version__ = "1.2.0"
__maintainer__ = "Tony Li"
__status__ = "R&D"

from office365.runtime.auth.authentication_context import AuthenticationContext
from office365.sharepoint.client_context import ClientContext
from office365.sharepoint.file import File
from office365.sharepoint.list import List
import pandas as pd
import os
from docx import Document, document
from openpyxl import load_workbook

# library for sharepoint
from pandas import DataFrame


def AuthenticateSharepoint(sharepoint_url, user, password):
    ctx_auth = AuthenticationContext(sharepoint_url)
    ctx_auth.acquire_token_for_user(user, password)
    ctx = ClientContext(sharepoint_url, ctx_auth)
    return ctx


def upload_file(context, local_path, web_path):
    with open(local_path, 'rb') as content_file:
        file_content = content_file.read()
        file_upload = bytes(file_content, 'utf-8')
        res = File.save_binary(context, web_path, file_upload)

    File.save_binary(context, web_path, file_content)
    print("File uploaded: {0}".webpath)


def download_file(context, web_path, local_path):
    response = File.open_binary(context, web_path)
    if response.ok:
        print_dash()
        with open(local_path, "wb") as local_file:
            local_file.write(response.content)
        print(web_path + '\r\ndownloaded at\r\n' + local_path)
        print_dash()


def print_dash():
    print('----------------------------')


def search_root(folders, files, context, tab=''):
    context.load(folders)
    context.execute_query()
    context.load(files)
    context.execute_query()
    return folders, files


def show_root(root, context):
    (folders, files) = search_root(root.folders, root.files, context)
    print('------------------------')
    for folder in folders:
        print("{0}".format(folder.properties["ServerRelativeUrl"]))
    for cur_file in files:
        print("{0}".format(cur_file.properties["ServerRelativeUrl"]))
    print('------------------------')
    return folders


# fcns to handle docx


def print_col(table, col_num):
    i = 1
    for cell in table.columns.__getitem__(col_num).cells:
        print(str(i) + '\t' + cell.text)
        i = i + 1


def col_2_list(table, col_num):
    ay = []
    for cell in table.columns.__getitem__(col_num).cells:
        ay.append(cell.text)
    return ay


def get_cell_text(table, col_num, cell_num):
    return table.columns.__getitem__(col_num).cells[cell_num].text


def process_BJEV_DTC_Doc(doc_name):
    doc = Document(doc_name)
    b = []
    for table in doc.tables:
        if get_cell_text(table, 0, 0) == '故障名称':
            b.append(col_2_list(table, 1))

    df: DataFrame = pd.DataFrame(b, columns=col_2_list(doc.tables[1], 0))
    return df


# Class to handle docx files, it can handle word files with card-view style tables
# mainly output dataframe and list object

def fetchTable(table):
    # grab the content from a docx class table to a composite list
    by = []

    for col in table.columns:
        ay = []
        for cell in col.cells:
            ay.append(cell.text)
        by.append(ay)

    return by


def appendTable2Lst(table, lst, force_fetch):
    # append the docx table content to a composite list

    def padLst(tb_header, ls_header, ls):
        # take content from (card-view like) ls to fill in a return r_ls
        # according to the match between ls_header and tb_header
        # tb_header is the baseline, ls_header is the datasource

        r_ls = [None] * len(tb_header)

        for hd_e in tb_header:
            if hd_e in ls_header:
                r_ls[tb_header.index(hd_e)] = ls[ls_header.index(hd_e)]

        return r_ls

    ls1 = fetchTable(table)

    if not ls1 or len(ls1) != 2:
        return lst

    if not lst:  # featch header from the first table
        lst.append(ls1[0])
        lst.append(ls1[1])
    elif (lst[0] == ls1[0]) or force_fetch:
        if len(lst[0]) == len(ls1[1]):
            # check if the table to be appended has the size and header
            lst.append(ls1[1])  # append content of the table
        else:
            lst.append(padLst(lst[0], ls1[0], ls1[1]))
            print("table do not match exactly")
    else:
        print('sth is wrong')
        pass

    return lst


class DocProcess:
    doc: document
    xls: pd.DataFrame
    file_path: str

    def __init__(self, file_name):
        self.file_path = file_name
        if 'xlsx' in file_name:
            self.xls = pd.ExcelFile(file_name)
            pass  # todo: implement excel handling
        if 'docx' in file_name:
            self.doc = Document(file_name)
        pass

    def export_Doc(self, file_path):
        if self.doc:
            self.doc.save(file_path)

    def export_Xls(self, file_path):
        if self.xls:
            self.xls.to_excel(file_path)

    def doc_tables_to_DF(self, force_fetch=True):
        ls = []
        for tb in self.doc.tables:
            ls = appendTable2Lst(tb, ls, force_fetch)
        return pd.DataFrame(ls[1:], columns=ls[0]), ls

    def compare_to_DF(self, internal_engine=True,
                      open_file=True):  # todo: add fcns to compare the doc to an external dataframe
        doc: document
        # the default output is card-like view
        # default output is a word file
        # os.startfile
        doc = Document()
        pass

    def update_file(self):  # right now only accepts xls file
        book = load_workbook(self.filepath)
        writer = pd.ExcelWriter(self.filepath, engine='openpyxl')
        writer.book = book
        writer.sheets = dict((ws.title, ws) for ws in book.worksheets)
        self.table.to_excel(writer, "Main", cols=['Diff1', 'Diff2'])
        writer.save()

    def process8Dfile(self):
        coversheet = pd.read_excel(self.xls, 'Coversheet')
        issue_id = coversheet.values[0][1]
        d_list = {}
        s_list = ('D0', 'D1', 'D2', 'D3', 'D4', 'D5',
                  'D6', 'D7', 'D8', 'Action')
        for it in s_list:
            if it == 'D0':
                tmp = pd.read_excel(self.xls, it, index_col=0, header=None).T
                tmp.rename(columns={tmp.columns[0]: "Identify"}, inplace=True)
            else:
                tmp = pd.read_excel(self.xls, it)
            d_list[it] = tmp
        d_list['D8'] = d_list['D8'].drop(columns=d_list['D8'].columns[0])

        return d_list, s_list

    def updateDF(self):  # todo: use the docs to update a dataframe
        pass
